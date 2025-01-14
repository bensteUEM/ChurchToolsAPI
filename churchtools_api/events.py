"""module containing parts used for events handling."""

import json
import logging
from datetime import datetime, timedelta
from pathlib import Path

import docx
import pytz
import requests

from churchtools_api.churchtools_api_abstract import ChurchToolsApiAbstract

logger = logging.getLogger(__name__)


class ChurchToolsApiEvents(ChurchToolsApiAbstract):
    """Part definition of ChurchToolsApi which focuses on events.

    Args:
        ChurchToolsApiAbstract: template with minimum references
    """

    def __init__(self) -> None:
        """Inherited initialization."""
        super()

    def get_events(self, **kwargs: dict) -> list[dict]:
        """Method to get all the events from given timespan or only the next event.

        Arguments:
            kwargs: optional params to modify the search criteria

        Keyword Arguments:
            eventId (int): number of event for single event lookup

            from_ (str|datetime): used as >= with starting date in format YYYY-MM-DD
                - added _ to name as opposed to ct_api because of reserved keyword
            to_ (str|datetime): used as < end date in format YYYY-MM-DD ONLY allowed
                with from_ - added _ to name as opposed to ct_api
                because of reserved keyword
            canceled (bool): If true, include also canceled events
            direction (str): direction of output 'forward' or 'backward'
                from the date defined by parameter 'from'
            limit (int): limits the number of events - Default = 1, if all events shall
                be retrieved insert 'None', only applies if direction is specified
            include (str): if Parameter is set to 'eventServices', the services of
                the event will be included

        Returns:
            list of events
        """
        url = self.domain + "/api/events"

        headers = {"accept": "application/json"}
        params = {"limit": 50}  # increases default pagination size

        if "eventId" in kwargs:
            url += "/{}".format(kwargs["eventId"])

        else:
            params = self._get_events_params_other(params=params, **kwargs)
            params = self._get_events_params_to_from(params=params, **kwargs)

        response = self.session.get(url=url, headers=headers, params=params)

        if response.status_code == requests.codes.ok:
            response_content = json.loads(response.content)
            response_data = self.combine_paginated_response_data(
                response_content,
                url=url,
                headers=headers,
                params=params,
            )
            return [response_data] if isinstance(response_data, dict) else response_data
        logger.warning(
            "%s Something went wrong fetching events: %s",
            response.status_code,
            response.content,
        )
        return None

    def _get_events_params_other(self, params: dict, **kwargs: dict) -> dict:
        """Helper function converting kwargs into params for request.

        Split in order to reduce function complexity
        converting everything except to and from

        Args:
            params: existing pre-set params
            kwargs: any additional options

        Returns:
            prepared params dict which can be used in request
        """
        if "canceled" in kwargs:
            params["canceled"] = kwargs["canceled"]
        if "direction" in kwargs:
            params["direction"] = kwargs["direction"]
        if "limit" in kwargs and "direction" in kwargs:
            params["limit"] = kwargs["limit"]
        elif "direction" in kwargs:
            logger.warning(
                "Use of limit is only allowed together with direction keyword",
            )
        if "include" in kwargs:
            params["include"] = kwargs["include"]
        return params

    def _get_events_params_to_from(self, params: dict, **kwargs: dict) -> dict:
        """Helper function converting kwargs into params for request.

        Split in order to reduce function complexity
        taking into account to and from only

        Args:
            params: existing pre-set params
            kwargs: any additional options

        Returns:
            prepared params dict which can be used in request
        """
        LENGTH_OF_DATE_WITH_HYPHEN = 10
        if "from_" in kwargs:
            from_ = kwargs["from_"]
            if isinstance(from_, datetime):
                from_ = from_.astimezone(pytz.utc).strftime("%Y-%m-%d")
            if len(from_) == LENGTH_OF_DATE_WITH_HYPHEN:
                params["from"] = from_
        if "to_" in kwargs and "from_" in kwargs:
            to_ = kwargs["to_"]
            if isinstance(to_, datetime):
                to_ = to_.astimezone(pytz.utc).strftime("%Y-%m-%d")
            if len(to_) == LENGTH_OF_DATE_WITH_HYPHEN:
                params["to"] = to_
        elif "to_" in kwargs:
            logger.warning("Use of to_ is only allowed together with from_")
        return params

    def get_event_by_calendar_appointment(
        self,
        appointment_id: int,
        start_date: str | datetime,
    ) -> dict:
        """This method is a helper to retrieve an event.

        for a specific calendar appointment including it's event services.

        Args:
            appointment_id: _description_
            start_date: either "2023-11-26T09:00:00Z", "2023-11-26" str or datetime

        Returns:
            event dict with event servics
        """
        if not isinstance(start_date, datetime):
            formats = {"iso": "%Y-%m-%dT%H:%M:%SZ", "date": "%Y-%m-%d"}
            for date_formats in formats.values():
                try:
                    start_date = datetime.strptime(start_date, date_formats).astimezone(
                        pytz.utc
                    )
                    break
                except ValueError:
                    continue

        events = self.get_events(
            from_=start_date,
            to_=start_date + timedelta(days=1),
            include="eventServices",
        )

        for event in events:
            if event["appointmentId"] == appointment_id:
                return event

        logger.info(
            "no event references appointment ID %s on start %s",
            appointment_id,
            start_date,
        )
        return None

    def get_AllEventData_ajax(self, eventId: int) -> dict:
        """Reverse engineered function from legacy AJAX API.

        which is used to get all event data for one event.
        Required to read special params not yet included in REST getEvents()
        Legacy AJAX request might stop working with any future release
            ... CSRF-Token is required in session header

        Arguments:
            eventId: number of the event to be requested

        Returns:
            event information
        """
        url = self.domain + "/index.php"
        headers = {"accept": "application/json"}
        params = {"q": "churchservice/ajax"}
        data = {"id": eventId, "func": "getAllEventData"}
        response = self.session.post(url=url, headers=headers, params=params, data=data)

        if response.status_code == requests.codes.ok:
            response_content = json.loads(response.content)
            if len(response_content["data"]) > 0:
                response_data = response_content["data"][str(eventId)]
                logger.debug("AJAX Event data len=%s", len(response_data))
                return response_data
            logger.info(
                "AJAX All Event data not successful - no event found:%s",
                response.status_code,
            )
            return None
        logger.info(
            "AJAX All Event data not successful: %s",
            response.status_code,
        )
        return None

    def get_event_services_counts_ajax(self, eventId: int, **kwargs: dict) -> dict:
        """Retrieve the number of services currently set for one specific event id.

        optionally get the number of services for one specific id on that event only.

        Arguments:
            eventId: id number of the calendar event
            **kwargs: keyword arguments as listed below

        Keywords:
            serviceId: id number of the service type to be filtered for
            serviceGroupId: id number of the group of services to request

        Returns:
            dict of service types and the number of services required for this event
        """
        event = self.get_events(eventId=eventId)[0]

        if "serviceId" in kwargs and "serviceGroupId" not in kwargs:
            service_count = 0
            for service in event["eventServices"]:
                if service["serviceId"] == kwargs["serviceId"]:
                    service_count += 1
            return {kwargs["serviceId"]: service_count}
        if "serviceId" not in kwargs and "serviceGroupId" in kwargs:
            all_services = self.get_services()
            serviceGroupServiceIds = [
                service["id"]
                for service in all_services
                if service["serviceGroupId"] == kwargs["serviceGroupId"]
            ]

            services = {}
            for service in event["eventServices"]:
                serviceId = service["serviceId"]
                if serviceId in serviceGroupServiceIds:
                    if serviceId in services:
                        services[serviceId] += 1
                    else:
                        services[serviceId] = 1

            return services
        logger.warning("Illegal combination of kwargs - check documentation either")
        return None

    def set_event_services_counts_ajax(
        self, eventId: int, serviceId: int, servicesCount: int
    ) -> bool:
        """Update the number of services currently set for one event specific id.

        Arguments:
            eventId: id number of the calendar event
            serviceId: id number of the service type to be filtered for
            servicesCount: number of services of the specified type to be planned
        Returns:
            successful execution
        """
        url = self.domain + "/index.php"
        headers = {"accept": "application/json"}
        params = {"q": "churchservice/ajax"}

        # restore other ServiceGroup assignments required for request form data

        services = self.get_services(returnAsDict=True)
        serviceGroupId = services[serviceId]["serviceGroupId"]
        servicesOfServiceGroup = self.get_event_services_counts_ajax(
            eventId,
            serviceGroupId=serviceGroupId,
        )
        # set new assignment
        servicesOfServiceGroup[serviceId] = servicesCount

        # Generate form specific data
        item_id = 0
        data = {"id": eventId, "func": "addOrRemoveServiceToEvent"}
        for item_id, (serviceIdRow, serviceCount) in enumerate(
            servicesOfServiceGroup.items()
        ):
            data[f"col{item_id}"] = serviceIdRow
            if serviceCount > 0:
                data[f"val{item_id}"] = "checked"
            data[f"count{item_id}"] = serviceCount

        response = self.session.post(url=url, headers=headers, params=params, data=data)

        if response.status_code == requests.codes.ok:
            response_content = json.loads(response.content)
            response_success = response_content["status"] == "success"

            number_match = (
                self.get_event_services_counts_ajax(eventId, serviceId=serviceId)[
                    serviceId
                ]
                == servicesCount
            )
            if number_match and response_success:
                return True
            logger.warning(
                "Request was successful but serviceId %s not changed to count %s ",
                serviceId,
                servicesCount,
            )
            return False
        logger.info(
            "set_event_services_counts_ajax not successful: %s",
            response.status_code,
        )
        return False

    def get_event_admins_ajax(self, eventId: int) -> list:
        """Get the admin id list of an event using legacy AJAX API.

        Params:
            eventId: number of the event to be changed
        Returns:
            list of admin ids.
        """
        event_data = self.get_AllEventData_ajax(eventId=eventId)
        if event_data is not None:
            if "admin" in event_data:
                admin_ids = [
                    int(available_event_id)
                    for available_event_id in event_data["admin"].split(",")
                ]
            else:
                admin_ids = []
            return admin_ids
        logger.info("No admins found because event not found")
        return []

    def set_event_admins_ajax(self, eventId: int, admin_ids: list) -> bool:
        """Set the admin id list of an event using legacy AJAX API.

        Parameters:
            eventId: number of the event to be changed
            admin_ids: list of admin user ids to be set as admin for event

        Returns:
            if successful.
        """
        url = self.domain + "/index.php"
        headers = {"accept": "application/json"}
        params = {"q": "churchservice/ajax"}

        data = {
            "id": eventId,
            "admin": ", ".join([str(admin_id) for admin_id in admin_ids]),
            "func": "updateEventInfo",
        }
        response = self.session.post(url=url, headers=headers, params=params, data=data)

        if response.status_code == requests.codes.ok:
            response_content = json.loads(response.content)
            response_data = response_content["status"] == "success"
            logger.debug(
                "Setting Admin IDs %s for event %s success", admin_ids, eventId
            )

            return response_data
        logger.info(
            "Setting Admin IDs %s for event %s failed with : %s",
            admin_ids,
            eventId,
            response.status_code,
        )
        return False

    def get_event_agenda(self, eventId: int) -> list:
        """Retrieve agenda for event by ID from ChurchTools.

        Arguments:
            eventId: number of the event
        Returns:
            list of event agenda items.
        """
        url = self.domain + f"/api/events/{eventId}/agenda"
        headers = {"accept": "application/json"}
        response = self.session.get(url=url, headers=headers)

        if response.status_code == requests.codes.ok:
            response_content = json.loads(response.content)
            response_data = response_content["data"].copy()
            logger.debug("Agenda load successful %s items", len(response_content))

            return response_data
        logger.info(
            "Event requested that does not have an agenda with status: %s",
            response.status_code,
        )
        return None

    def export_event_agenda(
        self, target_format: str, target_path: str = "./downloads", **kwargs: dict
    ) -> bool:
        """Exports the agenda as zip file for imports in presenter-programs.

        Parameters:
            target_format: fileformat or name of presentation software
                which should be supported.
                Supported formats are 'SONG_BEAMER', 'PRO_PRESENTER6'
                    and 'PRO_PRESENTER7'
            target_path: Filepath of the file which should
                be exported (including filename)
            kwargs: additional keywords as listed below

        Keywords:
            eventId: event id to check for agenda id should be exported
            agendaId: agenda id of the agenda which should be exported
                DO NOT combine with eventId because it will be overwritten!
            append_arrangement: if True, the name of the arrangement
                will be included within the agenda caption
            export_Songs: if True, the songfiles will be in the
                folder "Songs" within the zip file
            with_category: has no effect when exported in target format 'SONG_BEAMER'
        Returns:
            if successful.
        """
        if "eventId" in kwargs:
            if "agendaId" in kwargs:
                logger.warning(
                    "Invalid use of params - can not combine eventId and agendaId!",
                )
            else:
                agenda = self.get_event_agenda(eventId=kwargs["eventId"])
                agendaId = agenda["id"]
        elif "agendaId" in kwargs:
            agendaId = kwargs["agendaId"]
        else:
            logger.warning("Missing event or agendaId")
            return False

        # note: target path can be either a zip-file defined before function
        # call or just a folder
        is_zip = target_path.lower().endswith(".zip")
        if not is_zip:
            target_path = Path(target_path)
            target_path.mkdir(parents=True, exist_ok=True)

            if "eventId" in kwargs:
                new_file_name = "{}_{}.zip".format(agenda["name"], target_format)
            else:
                new_file_name = f"{target_format}_agendaId_{agendaId}.zip"

            target_path = target_path / new_file_name

        url = f"{self.domain}/api/agendas/{agendaId}/export"
        # NOTE the stream=True parameter below
        params = {"target": target_format}
        json_data = {}
        # The following 3 parameter 'appendArrangement', 'exportSongs' and
        # 'withCategory' are mandatory from the churchtools API side:
        json_data["appendArrangement"] = kwargs.get("append_arrangement", True)

        json_data["exportSongs"] = kwargs.get("export_songs", True)

        json_data["withCategory"] = kwargs.get("with_category", True)

        headers = {
            "accept": "application/json",
            "Content-Type": "application/json",
        }

        response = self.session.post(
            url=url,
            params=params,
            headers=headers,
            json=json_data,
        )
        result_ok = False
        if response.status_code == requests.codes.ok:
            response_content = json.loads(response.content)
            agenda_data = response_content["data"].copy()
            logger.debug("Agenda package found %s", response_content)
            result_ok = self.file_download_from_url(
                "{}/{}".format(self.domain, agenda_data["url"]),
                target_path,
            )
            if result_ok:
                logger.debug("download finished")
        else:
            logger.warning("export of event_agenda failed: %s", response.status_code)

        return result_ok

    def get_event_agenda_docx(self, agenda: dict, **kwargs: dict) -> docx.Document:
        """Generates custom docx document.

        Function to generate a custom docx document
        with the content of the event agenda from churchtools.

        Arguments:
            agenda: event agenda with services
            **kwargs: optional keywords as listed below

        Keywords:
            serviceGroups: list of servicegroup IDs that should be included
                - defaults to all if not supplied
            excludeBeforeEvent: bool: by default pre-event parts are excluded

        Returns:
            docx document reference
        """
        excludeBeforeEvent = kwargs.get("excludeBeforeEvent", False)

        logger.debug("Trying to get agenda for: %s", agenda["name"])

        document = docx.Document()
        heading = agenda["name"]
        heading += "- Draft" if not agenda["isFinal"] else ""
        document.add_heading(heading)
        modifiedDate = datetime.strptime(
            agenda["meta"]["modifiedDate"],
            "%Y-%m-%dT%H:%M:%S%z",
        )
        modifiedDate2 = modifiedDate.astimezone().strftime("%a %d.%m (%H:%M:%S)")
        document.add_paragraph(
            "Download from ChurchTools including changes until.: " + modifiedDate2,
        )

        agenda_item = 0  # Position Argument from Event Agenda is weird
        # therefore counting manually
        pre_event_last_item = True  # Event start is no item therefore look for change

        for item in agenda["items"]:
            if excludeBeforeEvent and item["isBeforeEvent"]:
                continue

            if item["type"] == "header":
                document.add_heading(item["title"], level=1)
                continue

            # helper for event start heading which is not part of the ct_api
            if pre_event_last_item and not item["isBeforeEvent"]:
                pre_event_last_item = False
                document.add_heading("Eventstart", level=1)

            agenda_item += 1

            title = str(agenda_item)
            title += " " + item["title"]

            if item["type"] == "song":
                title += ": " + item["song"]["title"]
                title += " (" + item["song"]["category"] + ")"

            document.add_heading(title, level=2)

            responsible_list = self._generate_responsible_list(item=item)
            responsible_text = ", ".join(responsible_list)
            document.add_paragraph(responsible_text)

            if item["note"] is not None and item["note"] != "":
                document.add_paragraph(item["note"])

            self._add_service_group_notes(
                document=document,
                service_group_notes=item["serviceGroupNotes"],
                service_groups=kwargs["serviceGroups"],
            )

        return document

    def _generate_responsible_list(self, item: dict) -> list:
        """Extracts information about the responsibility by agenda item.

        Args:
            item: the agenda item with all it's values

        Returns:
            prepared list of responsible persons used for further processing
        """
        responsible_list = []
        for responsible_item in item["responsible"]["persons"]:
            if responsible_item["person"] is not None:
                responsible_text = responsible_item["person"]["title"]
                if not responsible_item["accepted"]:
                    responsible_text += " (Angefragt)"
            else:
                responsible_text = "?"
            responsible_text += " " + responsible_item["service"] + ""
            responsible_list.append(responsible_text)

        if (
            len(item["responsible"]) > 0
            and len(item["responsible"]["persons"]) == 0
            and len(item["responsible"]["text"]) > 0
        ):
            responsible_list.append(
                item["responsible"]["text"]
                + " (Person statt Rolle in ChurchTools hinterlegt!)",
            )
        return responsible_list

    def _add_service_group_notes(
        self, document: docx.Document, service_group_notes: list, service_groups: dict
    ) -> None:
        """Subfunction which genereates service group note paragaphs.

        Args:
            document: the document item to work on
            service_group_notes: the list of items to consider
            service_groups: the service groups that are known
        """
        for item in service_group_notes:
            if len(item["serviceGroupNotes"]) > 0:
                for note in item["serviceGroupNotes"]:
                    if (
                        note["serviceGroupId"] in service_groups
                        and len(note["note"]) > 0
                    ):
                        document.add_heading(
                            "Bemerkung für {}:".format(
                                service_groups[note["serviceGroupId"]]["name"],
                            ),
                            level=4,
                        )
                        document.add_paragraph(note["note"])

    def get_persons_with_service(self, eventId: int, serviceId: int) -> list[dict]:
        """Helper function which should return the list of persons.

        that are assigned a specific service on a specific event.

        Args:
            eventId: id number from Events
            serviceId: id number from service masterdata

        Returns:
            list of persons
        """
        event = self.get_events(eventId=eventId)
        eventServices = event[0]["eventServices"]
        return [
            service for service in eventServices if service["serviceId"] == serviceId
        ]

    def get_event_masterdata(
        self, **kwargs: dict
    ) -> list | list[list] | dict | list[dict]:
        """Function to get the Masterdata of the event module.

        This information is required to map some IDs to specific items.

        Params
            kwargs: optional keywords as listed below

        Keywords:
            resultClass: str with name of the masterdata type (not datatype) common
             types are 'absenceReasons', 'songCategories', 'songSources'
            'services', 'serviceGroups', 'facts'

            returnAsDict: if the list with one type should be returned as dict by ID

        Returns:
            list of masterdata items, if multiple types list of lists (by type) if.
        """
        url = self.domain + "/api/event/masterdata"

        headers = {"accept": "application/json"}
        response = self.session.get(url=url, headers=headers)

        if response.status_code == requests.codes.ok:
            response_content = json.loads(response.content)
            response_data = response_content["data"].copy()

            if "resultClass" in kwargs:
                response_data = response_data[kwargs["resultClass"]]
                if kwargs.get("returnAsDict"):
                    response_data2 = response_data.copy()
                    response_data = {item["id"]: item for item in response_data2}
            logger.debug("Event Masterdata load successful len=%s", len(response_data))

            return response_data
        logger.info(
            "Event Masterdata requested failed: %s",
            response.status_code,
        )
        return None
